import os
import asyncio
import hashlib
import json
from typing import List, Dict, Any, Optional
from datetime import datetime
import logging
from pathlib import Path

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import google.generativeai as genai
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import PyPDF2
import docx
import io
from dotenv import load_dotenv
load_dotenv()

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Pydantic models
class Document(BaseModel):
    id: str
    name: str
    content: str
    chunks: List[str]
    embeddings: List[List[float]]
    created_at: str

class QueryRequest(BaseModel):
    query: str
    instruction: Optional[str] = None
    top_k: int = 5

class QueryResponse(BaseModel):
    answer: str
    relevant_chunks: List[str]
    sources: List[str]

class DocumentInfo(BaseModel):
    id: str
    name: str
    created_at: str
    chunks_count: int

# Knowledge Base class
class KnowledgeBase:
    def __init__(self):
        self.documents: Dict[str, Document] = {}
        self.data_file = "knowledge_base.json"
        self.load_data()
    
    def save_data(self):
        """Lưu dữ liệu vào file JSON"""
        try:
            data = {}
            for doc_id, doc in self.documents.items():
                data[doc_id] = {
                    "id": doc.id,
                    "name": doc.name,
                    "content": doc.content,
                    "chunks": doc.chunks,
                    "embeddings": doc.embeddings,
                    "created_at": doc.created_at
                }
            with open(self.data_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error(f"Error saving data: {e}")
    
    def load_data(self):
        """Load dữ liệu từ file JSON"""
        try:
            if os.path.exists(self.data_file):
                with open(self.data_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                for doc_id, doc_data in data.items():
                    self.documents[doc_id] = Document(**doc_data)
        except Exception as e:
            logger.error(f"Error loading data: {e}")
    
    def add_document(self, doc: Document):
        """Thêm tài liệu vào knowledge base"""
        self.documents[doc.id] = doc
        self.save_data()
    
    def remove_document(self, doc_id: str):
        """Xóa tài liệu khỏi knowledge base"""
        if doc_id in self.documents:
            del self.documents[doc_id]
            self.save_data()
            return True
        return False
    
    def get_documents(self) -> List[DocumentInfo]:
        """Lấy danh sách tài liệu"""
        return [
            DocumentInfo(
                id=doc.id,
                name=doc.name,
                created_at=doc.created_at,
                chunks_count=len(doc.chunks)
            )
            for doc in self.documents.values()
        ]
    
    def search_similar_chunks(self, query_embedding: List[float], top_k: int = 5):
        """Tìm kiếm chunks tương tự với query"""
        all_chunks = []
        all_embeddings = []
        chunk_sources = []
        
        for doc in self.documents.values():
            all_chunks.extend(doc.chunks)
            all_embeddings.extend(doc.embeddings)
            chunk_sources.extend([doc.name] * len(doc.chunks))
        
        if not all_embeddings:
            return [], []
        
        # Tính cosine similarity
        query_embedding = np.array(query_embedding).reshape(1, -1)
        embeddings_matrix = np.array(all_embeddings)
        
        similarities = cosine_similarity(query_embedding, embeddings_matrix)[0]
        
        # Lấy top_k chunks có similarity cao nhất
        top_indices = np.argsort(similarities)[::-1][:top_k]
        
        relevant_chunks = [all_chunks[i] for i in top_indices]
        sources = [chunk_sources[i] for i in top_indices]
        
        return relevant_chunks, sources

# Initialize components
app = FastAPI(title="RAG System Server", version="1.0.0")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Gemini
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=GEMINI_API_KEY)

# Initialize models
generation_model = genai.GenerativeModel('gemini-2.5-flash-lite')

# Initialize Knowledge Base
kb = KnowledgeBase()

# Helper functions
def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """Trích xuất text từ file"""
    try:
        if filename.endswith('.pdf'):
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        
        elif filename.endswith('.docx'):
            doc_file = io.BytesIO(file_content)
            doc = docx.Document(doc_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        
        elif filename.endswith('.txt'):
            return file_content.decode('utf-8')
        
        else:
            raise ValueError(f"Unsupported file format: {filename}")
            
    except Exception as e:
        logger.error(f"Error extracting text from {filename}: {e}")
        raise HTTPException(status_code=400, detail=f"Error processing file: {str(e)}")

def split_text_into_chunks(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Chia text thành các chunks"""
    words = text.split()
    chunks = []
    
    for i in range(0, len(words), chunk_size - overlap):
        chunk = ' '.join(words[i:i + chunk_size])
        if chunk.strip():
            chunks.append(chunk.strip())
    
    return chunks

async def get_embeddings(texts: List[str]) -> List[List[float]]:
    """Lấy embeddings cho danh sách texts"""
    embeddings = []
    for text in texts:
        try:
            result = await asyncio.to_thread(
                genai.embed_content,
                model="models/text-embedding-004",
                content=text,
                task_type="retrieval_document"
            )
            embeddings.append(result['embedding'])
        except Exception as e:
            logger.error(f"Error getting embedding: {e}")
            # Fallback: zero vector
            embeddings.append([0.0] * 768)
    
    return embeddings

async def get_query_embedding(query: str) -> List[float]:
    """Lấy embedding cho query"""
    try:
        result = await asyncio.to_thread(
            genai.embed_content,
            model="models/text-embedding-004",
            content=query,
            task_type="retrieval_query"
        )
        return result['embedding']
    except Exception as e:
        logger.error(f"Error getting query embedding: {e}")
        return [0.0] * 768

def generate_document_id(content: str) -> str:
    """Tạo ID unique cho document"""
    return hashlib.md5(content.encode()).hexdigest()

# API Endpoints
@app.post("/documents/")
async def add_document(file: UploadFile = File(...), request: Request = None):
    """Thêm tài liệu vào Knowledge Base"""
    try:
        # Lấy API key từ header nếu có
        api_key = request.headers.get("X-GEMINI-API-KEY") if request else None
        if api_key:
            genai.configure(api_key=api_key)

        # Đọc nội dung file
        file_content = await file.read()
        
        # Trích xuất text
        text_content = extract_text_from_file(file_content, file.filename)
        
        # Tạo document ID
        doc_id = generate_document_id(text_content)
        
        # Kiểm tra xem document đã tồn tại chưa
        if doc_id in kb.documents:
            raise HTTPException(status_code=409, detail="Document already exists")
        
        # Chia text thành chunks
        chunks = split_text_into_chunks(text_content)
        
        if not chunks:
            raise HTTPException(status_code=400, detail="No valid text chunks found in document")
        
        # Lấy embeddings cho chunks
        embeddings = await get_embeddings(chunks)
        
        # Tạo document object
        document = Document(
            id=doc_id,
            name=file.filename,
            content=text_content,
            chunks=chunks,
            embeddings=embeddings,
            created_at=datetime.now().isoformat()
        )
        
        # Thêm vào knowledge base
        kb.add_document(document)
        
        return {
            "message": "Document added successfully",
            "document_id": doc_id,
            "chunks_count": len(chunks)
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error adding document: {e}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")

@app.delete("/documents/{document_id}")
async def delete_document(document_id: str):
    """Xóa tài liệu khỏi Knowledge Base"""
    if kb.remove_document(document_id):
        return {"message": "Document deleted successfully"}
    else:
        raise HTTPException(status_code=404, detail="Document not found")

@app.get("/documents/")
async def list_documents() -> List[DocumentInfo]:
    """Lấy danh sách tài liệu"""
    return kb.get_documents()

@app.post("/query/")
async def query_knowledge_base(request: QueryRequest, fastapi_request: Request) -> QueryResponse:
    """Query vào Knowledge Base"""
    try:
        if not kb.documents:
            raise HTTPException(status_code=400, detail="No documents in knowledge base")
        
        # Lấy API key từ header nếu có
        api_key = fastapi_request.headers.get("X-GEMINI-API-KEY")
        if api_key:
            genai.configure(api_key=api_key)

        # Lấy embedding cho query
        query_embedding = await get_query_embedding(request.query)
        
        # Tìm kiếm chunks tương tự
        relevant_chunks, sources = kb.search_similar_chunks(query_embedding, request.top_k)
        
        if not relevant_chunks:
            return QueryResponse(
                answer="Không tìm thấy thông tin liên quan đến câu hỏi của bạn.",
                relevant_chunks=[],
                sources=[]
            )
        
        # Tạo context từ relevant chunks
        context = "\n\n".join(relevant_chunks)
        
        # Tạo instruction mặc định hoặc sử dụng instruction từ user
        default_instruction = """Bạn là một AI assistant thông minh. Hãy trả lời câu hỏi dựa trên thông tin được cung cấp trong context. 
        Nếu thông tin không đủ để trả lời, hãy nói rõ điều đó. Trả lời bằng tiếng Việt một cách chi tiết và dễ hiểu.
        Lưu ý chỉ tập trung rõ vào query được hỏi, không được trả lời thông tin không liên quan tới query, không cần mở rộng kiểu bạn cần thêm này kia hay không.
        Các tên riêng, tên dự án,... cần được chú ý rõ, phải đúng chuẩn 100% tên như người dùng hỏi mới trả lời, không trả lời những thứ gần giống.
        Ưu tiên trả lời ngắn gọn, súc tích.
        """
        
        instruction = request.instruction if request.instruction else default_instruction
        
        # Tạo prompt cho Gemini
        prompt = f"""
{instruction}

Context:
{context}

Câu hỏi: {request.query}

Trả lời:
"""
        
        # Generate response
        response = await asyncio.to_thread(generation_model.generate_content, prompt)
        
        return QueryResponse(
            answer=response.text,
            relevant_chunks=relevant_chunks,
            sources=sources
        )
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error querying knowledge base: {e}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")

@app.get("/")
async def root():
    """Health check endpoint"""
    return {
        "message": "RAG System Server is running",
        "documents_count": len(kb.documents),
        "total_chunks": sum(len(doc.chunks) for doc in kb.documents.values())
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)